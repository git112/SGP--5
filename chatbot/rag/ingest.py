import os
import re
import time
from collections import defaultdict
from typing import List
from dotenv import load_dotenv
from docx import Document as DocxDocument
from openpyxl import load_workbook
from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_nomic import NomicEmbeddings
from langchain_pinecone import PineconeVectorStore
from pinecone import Pinecone, ServerlessSpec

load_dotenv()

DATA_DIR            = "./data"
JD_DIR              = os.path.join(DATA_DIR, "Job description")
PINECONE_INDEX_NAME = "placementor"
NAMESPACE           = "placements"
BATCH_SIZE          = 100

text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=800,
    chunk_overlap=120,
    separators=[". ", "! ", "? ", "\n\n", "\n", " "]
)

# ── Helpers ───────────────────────────────────────────────────────────────────
def extract_year(filename: str) -> str:
    match = re.search(r'20\d\d', filename)
    return match.group() if match else ""

def clean_val(val) -> str:
    if val is None:
        return ""
    s = str(val).strip()
    return "" if s.startswith("=") else s

def is_empty_row(row) -> bool:
    return all(clean_val(c.value) == "" for c in row)

def is_header_row(row) -> bool:
    values = [clean_val(c.value) for c in row if clean_val(c.value)]
    if not values:
        return False
    keywords = ['id', 'company', 'package', 'gender', 'cgpa', 'placed',
                'role', 'year', 'sr', 'name', 'offer', 'location', 'city',
                'remarks', 'placement', 'policy']
    matches = sum(1 for v in values if any(k in v.lower() for k in keywords))
    return matches >= 2

def find_header_row(ws) -> int:
    """Returns 1-based index of the actual header row."""
    for i, row in enumerate(ws.iter_rows(min_row=1, max_row=10), start=1):
        if is_empty_row(row):
            continue
        raw = [str(c.value) for c in row if c.value is not None]
        if raw and all(v.startswith("=") for v in raw):
            continue
        if is_header_row(row):
            return i
    return 1

# ── PDF Loader ────────────────────────────────────────────────────────────────
def load_pdfs() -> List[Document]:
    documents = []
    if not os.path.exists(JD_DIR):
        print(f"JD directory not found: {JD_DIR}")
        return []

    for file in os.listdir(JD_DIR):
        if not file.endswith(".pdf"):
            continue
        path = os.path.join(JD_DIR, file)
        try:
            loader = PyPDFLoader(path)
            pages = loader.load()
            for doc in pages:
                doc.metadata.update({
                    "source": file,
                    "type": "job_description",
                    "year": ""
                })
            documents.extend(pages)
            print(f"  ✅ PDF: {file} ({len(pages)} pages)")
        except Exception as e:
            print(f"  ❌ PDF Error {file}: {e}")
    return documents

# ── DOCX Loader ───────────────────────────────────────────────────────────────
def load_docx() -> List[Document]:
    documents = []
    if not os.path.exists(JD_DIR):
        return []

    for file in os.listdir(JD_DIR):
        if not file.endswith(".docx"):
            continue
        path = os.path.join(JD_DIR, file)
        try:
            doc = DocxDocument(path)
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            if text:
                documents.append(Document(
                    page_content=text,
                    metadata={"source": file, "type": "job_description", "year": ""}
                ))
                print(f"  ✅ DOCX: {file}")
        except Exception as e:
            print(f"  ❌ DOCX Error {file}: {e}")
    return documents

# ── Smart Excel Loader ────────────────────────────────────────────────────────
def load_excel_files() -> tuple:
    """
    Loads all Excel files with smart header detection.
    Handles:
    - Normal files: Row 1 = headers
    - 2025.xlsx: Row 1 = formulas, Row 2 = headers, Row 3 = data
    - Empty files: 2024-FINAL, 2025--- (skipped gracefully)
    Returns (documents, raw_rows_for_summaries)
    """
    documents = []
    all_raw_rows = []

    for file in sorted(os.listdir(DATA_DIR)):
        if not file.endswith(".xlsx") or file.startswith("~$"):
            continue

        filepath = os.path.join(DATA_DIR, file)
        year = extract_year(file)
        is_insights = "Insights" in file
        doc_type = "placement_insight" if is_insights else "placement_record"

        try:
            wb = load_workbook(filepath, data_only=True)

            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]

                # Find real header row
                header_row_idx = find_header_row(ws)
                header_cells = list(
                    ws.iter_rows(min_row=header_row_idx, max_row=header_row_idx)
                )[0]
                headers = [
                    clean_val(c.value) or f"Col{i+1}"
                    for i, c in enumerate(header_cells)
                ]

                # Trim trailing Col headers
                while headers and re.match(r'^Col\d+$', headers[-1]):
                    headers.pop()

                if not headers:
                    continue

                row_count = 0
                for row in ws.iter_rows(min_row=header_row_idx + 1):
                    if is_empty_row(row):
                        continue

                    row_data = {}
                    for i, cell in enumerate(row):
                        if i >= len(headers):
                            break
                        val = clean_val(cell.value)
                        if val:
                            row_data[headers[i]] = val

                    if len(row_data) < 2:
                        continue

                    # Build readable text
                    parts = []
                    for key, val in row_data.items():
                        if not key or key.startswith("Col"):
                            continue
                        # Normalize package
                        if "package" in key.lower():
                            try:
                                pkg_str = val.replace("LPA", "").replace("lpa", "").strip()
                                if "to" in pkg_str:
                                    lo, hi = pkg_str.split("to")
                                    val = f"{lo.strip()} to {hi.strip()} LPA"
                                else:
                                    num = float(pkg_str.split()[0])
                                    val = f"{num} LPA"
                            except:
                                if "lpa" not in val.lower():
                                    val = f"{val} LPA"
                        # Normalize gender
                        if "gender" in key.lower():
                            val = val.title()
                        parts.append(f"{key.strip()}: {val}")

                    if not parts:
                        continue

                    prefix = f"Year: {year} | " if year else ""
                    content = prefix + " | ".join(parts)

                    documents.append(Document(
                        page_content=content,
                        metadata={
                            "source": file,
                            "sheet": sheet_name,
                            "type": doc_type,
                            "year": year
                        }
                    ))

                    if not is_insights:
                        all_raw_rows.append({**row_data, "year": year})

                    row_count += 1

                if row_count > 0:
                    print(f"  ✅ {file} | {sheet_name} | {row_count} rows | Headers: {headers[:4]}")
                else:
                    print(f"  ⚠️  {file} | {sheet_name} | empty or unreadable")

        except Exception as e:
            print(f"  ❌ Excel Error {file}: {e}")

    return documents, all_raw_rows

# ── Summary Generator ─────────────────────────────────────────────────────────
def generate_summaries(all_raw_rows: List[dict]) -> List[Document]:
    """
    Creates year + company summary documents.
    Enables accurate answers to:
    - 'which company hired most in 2023'
    - 'highest package in 2024'
    - 'how many students placed in 2022'
    """
    docs = []
    by_year = defaultdict(list)

    for row in all_raw_rows:
        year = row.get("year", "")
        if year:
            by_year[year].append(row)

    for year, rows in sorted(by_year.items()):
        company_counts   = defaultdict(int)
        company_packages = defaultdict(list)
        company_roles    = defaultdict(set)

        for row in rows:
            company = ""
            for key in row:
                if "company" in key.lower() and row[key] not in ["None", "nan", ""]:
                    company = str(row[key]).strip()
                    break
            if not company:
                continue

            company_counts[company] += 1

            for key in row:
                if "package" in key.lower() and row[key]:
                    try:
                        pkg_str = str(row[key]).replace("LPA", "").replace("lpa", "").strip()
                        if "to" in pkg_str:
                            parts = pkg_str.split("to")
                            num = (float(parts[0].strip()) + float(parts[1].strip())) / 2
                        else:
                            num = float(pkg_str.split()[0])
                        company_packages[company].append(num)
                    except:
                        pass

            for key in row:
                if any(k in key.lower() for k in ["role", "remark", "position"]):
                    val = str(row[key]).strip()
                    if val and val not in ["None", "nan", "1", "0", ""]:
                        company_roles[company].add(val)

        if not company_counts:
            continue

        total   = sum(company_counts.values())
        top10   = sorted(company_counts.items(), key=lambda x: x[1], reverse=True)[:10]
        top_str = ", ".join([f"{c} ({n} students)" for c, n in top10])

        all_pkgs = [p for pkgs in company_packages.values() for p in pkgs]
        highest  = max(all_pkgs) if all_pkgs else 0
        average  = sum(all_pkgs) / len(all_pkgs) if all_pkgs else 0
        highest_company = next(
            (c for c, pkgs in company_packages.items() if pkgs and max(pkgs) == highest),
            "Unknown"
        )

        # Year summary
        docs.append(Document(
            page_content=f"""YEAR {year} PLACEMENT SUMMARY:
Total students placed in {year}: {total}
Total companies that visited in {year}: {len(company_counts)}
Top hiring companies in {year}: {top_str}
All companies that hired in {year}: {', '.join(sorted(company_counts.keys()))}
Highest package in {year}: {highest} LPA (offered by {highest_company})
Average package in {year}: {average:.2f} LPA
""",
            metadata={"source": f"Summary_{year}", "type": "year_summary", "year": year}
        ))

        # Per-company summaries
        for company, count in company_counts.items():
            pkgs  = company_packages.get(company, [])
            roles = company_roles.get(company, set())
            avg_c = sum(pkgs) / len(pkgs) if pkgs else 0
            max_c = max(pkgs) if pkgs else 0

            docs.append(Document(
                page_content=f"""COMPANY SUMMARY: {company} | YEAR: {year}
{company} hired {count} students in {year}
Package offered by {company} in {year}: {max_c} LPA{f' (avg: {avg_c:.2f} LPA)' if len(pkgs) > 1 else ''}
Roles offered by {company}: {', '.join(roles) if roles else 'Not specified'}
""",
                metadata={
                    "source": f"CompanySummary_{year}",
                    "type": "company_summary",
                    "year": year,
                    "company": company
                }
            ))

    print(f"  ✅ Generated {len(docs)} summary documents")
    return docs

# ── Chunking ──────────────────────────────────────────────────────────────────
def chunk_documents(all_docs: List[Document]) -> List[Document]:
    final_chunks = []

    for doc in all_docs:
        doc_type = doc.metadata.get("type", "")

        # Keep short structured docs as single chunks
        if doc_type in ["placement_record", "placement_insight",
                        "year_summary", "company_summary"] and len(doc.page_content) < 400:
            doc.page_content = f"{doc.metadata.get('source', '')}. {doc.page_content}"
            final_chunks.append(doc)
        else:
            chunks = text_splitter.split_documents([doc])
            for chunk in chunks:
                source = chunk.metadata.get("source", "Placement Data")
                chunk.page_content = f"{source}. {chunk.page_content}"
                final_chunks.append(chunk)

    return final_chunks

# ── Pinecone Upload ───────────────────────────────────────────────────────────
def upload_to_pinecone(chunks: List[Document]):
    embeddings = NomicEmbeddings(model="nomic-embed-text-v1.5")
    pc = Pinecone(api_key=os.getenv("PINECONE_API_KEY"))

    existing_names = pc.list_indexes().names()

    if PINECONE_INDEX_NAME not in existing_names:
        print(f"Creating Pinecone index: {PINECONE_INDEX_NAME}")
        pc.create_index(
            name=PINECONE_INDEX_NAME,
            dimension=768,
            metric="cosine",
            spec=ServerlessSpec(cloud="aws", region="us-east-1")
        )
        while not pc.describe_index(PINECONE_INDEX_NAME).status['ready']:
            time.sleep(1)
    else:
        print(f"Clearing old vectors...")
        pc.Index(PINECONE_INDEX_NAME).delete(delete_all=True, namespace=NAMESPACE)
        time.sleep(3)

    total_batches = (len(chunks) - 1) // BATCH_SIZE + 1
    print(f"Uploading {len(chunks)} chunks in {total_batches} batches...")

    for i in range(0, len(chunks), BATCH_SIZE):
        batch = chunks[i:i + BATCH_SIZE]
        batch_num = i // BATCH_SIZE + 1
        try:
            PineconeVectorStore.from_documents(
                batch,
                embeddings,
                index_name=PINECONE_INDEX_NAME,
                namespace=NAMESPACE
            )
            print(f"  Uploaded batch {batch_num}/{total_batches}")
        except Exception as e:
            print(f"  ❌ Batch {batch_num} failed: {e}")

# ── Main ──────────────────────────────────────────────────────────────────────
def ingest():
    print("=" * 60)
    print("PlaceMentor AI — Data Ingestion")
    print("=" * 60)

    print("\n📄 Loading PDFs...")
    pdf_docs = load_pdfs()

    print("\n📝 Loading DOCX...")
    docx_docs = load_docx()

    print("\n📊 Loading Excel files...")
    excel_docs, raw_rows = load_excel_files()

    print("\n📈 Generating summaries...")
    summary_docs = generate_summaries(raw_rows)

    all_docs = pdf_docs + docx_docs + excel_docs + summary_docs

    print(f"\n✂️  Chunking {len(all_docs)} documents...")
    final_chunks = chunk_documents(all_docs)

    type_counts = defaultdict(int)
    for doc in final_chunks:
        type_counts[doc.metadata.get("type", "unknown")] += 1

    print(f"\n{'=' * 60}")
    print(f"TOTAL CHUNKS TO UPLOAD: {len(final_chunks)}")
    for t, c in sorted(type_counts.items()):
        print(f"  {t}: {c}")
    print("=" * 60)

    print("\n🚀 Uploading to Pinecone...")
    upload_to_pinecone(final_chunks)

    print("\n✅ Ingestion complete!")

if __name__ == "__main__":
    ingest()